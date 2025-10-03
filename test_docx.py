#!/usr/bin/env python3
"""Teste simples para verificar se a geração DOCX funciona"""

try:
    from docx import Document
    from docx.shared import Inches
    import matplotlib.pyplot as plt
    import io
    import base64
    print("✅ Módulos importados com sucesso!")
    
    # Teste básico de criação de documento
    document = Document()
    document.add_heading('Teste DOCX', 0)
    document.add_paragraph('Este é um teste de geração de documento Word.')
    
    # Teste de gráfico
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.bar(['A', 'B', 'C'], [1, 2, 3])
    ax.set_title('Gráfico de Teste')
    
    # Salva gráfico em memória
    img_bytes = io.BytesIO()
    fig.savefig(img_bytes, format='png')
    plt.close(fig)
    img_bytes.seek(0)
    
    # Adiciona gráfico ao documento
    document.add_picture(img_bytes, width=Inches(5.5))
    
    # Salva documento em memória
    doc_bytes = io.BytesIO()
    document.save(doc_bytes)
    doc_bytes.seek(0)
    
    # Converte para base64
    b64_content = base64.b64encode(doc_bytes.read()).decode('utf-8')
    
    print(f"✅ DOCX gerado com sucesso! Tamanho: {len(b64_content)} caracteres")
    print("✅ Teste concluído com sucesso!")
    
except Exception as e:
    print(f"❌ Erro no teste: {e}")
    import traceback
    traceback.print_exc()

